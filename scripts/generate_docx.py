#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""根据分析结果 JSON，生成包含明细表格和附件凭证页的 DOCX 报销单。"""

import sys
import os
import re
import json
import math
import argparse
import xml.etree.ElementTree as ET
from datetime import datetime

# 确保 python-docx 已安装
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：未安装 python-docx，请执行：pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 1.5
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM  # 18cm
MAX_IMG_WIDTH_CM = CONTENT_WIDTH_CM
MAX_IMG_HEIGHT_CM = PAGE_HEIGHT_CM - 2 * MARGIN_CM  # 26.7cm，但取 24cm 保留空间

GROUP_ORDER = ["flight", "train", "ride_hailing", "self_drive", "hotel", "other", "meal_allowance", "self_drive_subsidy"]
GROUP_CN = {
    "flight": "飞机", "train": "高铁", "ride_hailing": "网约车",
    "self_drive": "自驾", "hotel": "住宿", "other": "其它",
    "meal_allowance": "误餐补贴", "self_drive_subsidy": "自驾补贴",
}
WEEKDAY_CN = ["一", "二", "三", "四", "五", "六", "日"]


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------
def get_weekday_cn(date_str):
    """根据日期字符串返回中文星期几。"""
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    return WEEKDAY_CN[dt.weekday()]


def format_amount(amount):
    """格式化金额为千分位两位小数，如 2,707.80。"""
    return f"{amount:,.2f}"


def set_cell_border(cell):
    """为单元格设置上下左右细边框（1pt）。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_run_font(run, font_name="宋体", size=None, bold=False):
    """为 run 设置字体，同时锁定 ASCII / 中文(含 cs) 字体，保证 WPS 与 Office 渲染一致。"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    )
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)
    if size is not None:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True


def lock_document_font(doc, font_name="宋体", size=12):
    """锁定整篇文档默认字体（Normal 样式），避免 WPS / Office 渲染差异。"""
    try:
        style = doc.styles["Normal"]
    except KeyError:
        return
    style.font.name = font_name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    )
    old = rpr.find(qn("w:rFonts"))
    if old is not None:
        rpr.remove(old)
    rpr.append(rfonts)


def display_width(text):
    """估算文本显示宽度（CJK 记 2，其余记 1）。"""
    w = 0
    for ch in str(text):
        w += 2 if ord(ch) > 0x2E80 else 1
    return w


def compute_col_widths(matrix, total_cm=17.5, min_w=None):
    """根据各列实际内容多寡计算列宽（cm）；matrix 第一行通常为表头。"""
    ncols = len(matrix[0])
    col_units = [0] * ncols
    for row in matrix:
        for i, c in enumerate(row):
            col_units[i] = max(col_units[i], display_width(c))
    total_u = sum(col_units) or 1
    widths = [col_units[i] / total_u * total_cm for i in range(ncols)]
    if min_w:
        widths = [max(widths[i], min_w[i]) for i in range(ncols)]
    s = sum(widths)
    if s > 0 and abs(s - total_cm) > 0.01:
        widths = [w * total_cm / s for w in widths]
    return widths


def set_cell_width(cell, cm):
    """设置单元格宽度（twips），配合固定布局锁定列宽。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")}/>')
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tcW.set(qn("w:type"), "dxa")


def enable_word_wrap(cell):
    """开启单元格自动换行，内容完整展示不溢出。"""
    for p in cell.paragraphs:
        pPr = p._p.get_or_add_pPr()
        ww = pPr.find(qn("w:wordWrap"))
        if ww is None:
            ww = parse_xml(f'<w:wordWrap {nsdecls("w")} w:val="on"/>')
            pPr.append(ww)


def _sort_key(exp):
    """返回可排序的日期键（YYYY-MM-DD），用于明细表按发生时间排序。"""
    sk = exp.get("sort_key")
    if sk:
        return sk
    d = exp.get("date_display") or exp.get("date") or ""
    m = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', d)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    m = re.search(r'(\d{1,2})月(\d{1,2})日', d)
    if m:
        return f"2026-{int(m.group(1)):02d}-{int(m.group(2)):02d}"
    return "9999-99-99"


def make_paragraph(doc, text, font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before=0, space_after=0, font_name="宋体"):
    """创建段落并设置字体。返回 paragraph 对象。"""
    para = doc.add_paragraph()
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = para.add_run(text)
    set_run_font(run, font_name=font_name, size=font_size, bold=bold)
    return para


def set_row_height(row, height_cm):
    """设置表格行高（近似）。"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm / 2.54 * 72 * 20)}" w:hRule="exact"/>'
    )
    trPr.append(trHeight)


def add_shading_to_paragraph(para, color_hex):
    """为段落添加底纹（用于标注行背景色）。"""
    pPr = para._element.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    pPr.append(shd)


def set_cell_vertical_alignment(cell, align="center"):
    """设置单元格垂直对齐。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign_map = {"center": "center", "top": "top", "bottom": "bottom"}
    vAlign = parse_xml(
        f'<w:vAlign {nsdecls("w")} w:val="{vAlign_map.get(align, "center")}"/>'
    )
    tcPr.append(vAlign)


def image_natural_cm(path):
    """返回图片在自然 96dpi 下的 (宽cm, 高cm, 宽高比)，用于保持原始比例不变形。"""
    from PIL import Image
    img = Image.open(path)
    w_px, h_px = img.size
    dpi = 96.0
    w_cm = w_px / dpi * 2.54
    h_cm = h_px / dpi * 2.54
    ratio = w_cm / h_cm if h_cm > 0 else 1.0
    return w_cm, h_cm, ratio


def is_invoice_image(path):
    """判断图片是否为发票（发票保持最宽显示）。"""
    name = os.path.basename(path).lower()
    keywords = ["发票", "invoice"]
    return any(k in name for k in keywords)


def is_doc_voucher(path):
    """判断图片是否为文档型凭证（PDF 转出的发票/行程单/报销单/账单等），
    应占满整行、与页面同宽显示。"""
    name = os.path.basename(path).lower()
    keywords = ["发票", "invoice", "行程单", "报销单", "账单", "凭证",
                "itinerary", "boarding"]
    return any(k in name for k in keywords)


def layout_images_continuous(doc, image_paths, state):
    """
    连续拼版：跨类别连续排版（不强制每类分页），仅在整行放不下时才换页；
    图片始终按原始宽高比显示——只指定宽度，高度由 python-docx 自动计算，
    绝不拉伸变形。
    - 发票 / 宽图（宽高比>=1.1）：占满整行（18cm 宽）
    - 窄图（截图等）：与其它窄图并排，每行最多 2 张
    state['page_used_h'] 累计当前页已用高度，跨多次调用保持连续。
    """
    if not image_paths:
        return

    MAX_W_CM = MAX_IMG_WIDTH_CM          # 18
    MAX_H_CM = 26.7                      # 单页内容区最大高度（A4 29.7 - 边距2×1.5 = 26.7）
    GAP_CM = 0.25
    MIN_NARROW_W = 6.0
    HALF_W = MAX_W_CM / 2 - GAP_CM / 2   # 并排时单列最大宽度

    # 预计算每张图的显示尺寸（严格保持原始宽高比）
    items = []
    for path in image_paths:
        if not os.path.exists(path):
            continue
        try:
            w_cm, h_cm, ratio = image_natural_cm(path)
        except Exception:
            continue
        is_inv = is_invoice_image(path)
        full = is_doc_voucher(path) or ratio >= 1.1
        if full:
            # 文档凭证/宽图：目标 18cm 但不超自然宽度
            dw = min(MAX_W_CM, w_cm)
            if dw / ratio <= MAX_H_CM:
                dh = dw / ratio
            else:
                dh, dw = MAX_H_CM, MAX_H_CM * ratio
        else:
            # 窄图（截图等）：统一用半行宽，不超自然宽度 → 绝不放大变模糊
            dw = min(w_cm, HALF_W)
            dh = dw / ratio
            dw = max(dw, MIN_NARROW_W)
        items.append({"path": path, "dw": dw, "dh": dh, "full": full})

    def flush(buf):
        if not buf:
            return
        row_h = max(it["dh"] for it in buf)
        if state["page_used_h"] + row_h > MAX_H_CM:
            doc.add_page_break()
            state["page_used_h"] = 0.0
        if len(buf) == 1:
            it = buf[0]
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run()
            run.add_picture(it["path"], width=Cm(it["dw"]))   # 仅指定宽度→高度自动，不变形
            if not it["full"]:
                # 窄图独行：上下均匀留白，视觉平衡替代放大填满
                free_h = MAX_H_CM - row_h
                if free_h > 2.0:
                    half = free_h / 2
                    para.paragraph_format.space_before = Cm(half)
                    para.paragraph_format.space_after = Cm(half)
                    # 页面已满，重置累计高度让下一行自然换页
                    state["page_used_h"] = MAX_H_CM
                else:
                    state["page_used_h"] += row_h + GAP_CM
            else:
                state["page_used_h"] += row_h + GAP_CM
            state["page_used_h"] += row_h + GAP_CM
        else:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_xml = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tblBorders>'
            )
            tbl._tbl.tblPr.append(tbl_xml)
            if state["page_used_h"] + row_h > MAX_H_CM:
                doc.add_page_break()
                state["page_used_h"] = 0.0
            for ci, it in enumerate(buf):
                cell = tbl.rows[0].cells[ci]
                cell.width = Cm(it["dw"])
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                run = para.add_run()
                run.add_picture(it["path"], width=Cm(it["dw"]))
            state["page_used_h"] += row_h + GAP_CM

    # 连续填充：窄图累积到 2 张成一行；满行图单独成行
    buf = []
    for it in items:
        if it["full"]:
            flush(buf)
            buf = []
            flush([it])
        else:
            buf.append(it)
            if len(buf) >= 2:
                flush(buf)
                buf = []
    flush(buf)


# ---------------------------------------------------------------------------
# 清除 DOCX 中的 python-docx 元数据痕迹
# ---------------------------------------------------------------------------
def _post_process_docx(docx_path):
    """后处理：用 python-docx 清空页眉页脚，再用 ZIP 操作清除 rsid 和元数据。"""
    import zipfile, shutil
    from lxml import etree

    # Step 1: 用 python-docx 清空页眉页脚并移除引用
    from docx import Document as _Doc
    d = _Doc(docx_path)
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for section in d.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer]:
            if hf is not None:
                for child in list(hf._element):
                    hf._element.remove(child)
        sectPr = section._sectPr
        for tag_name in ["headerReference", "footerReference"]:
            for el in sectPr.findall("{%s}%s" % (NS, tag_name)):
                sectPr.remove(el)
    d.save(docx_path)

    # Step 2: ZIP 级清除 rsid 和元数据（不删除任何文件，避免悬空引用）
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    raw = data.decode("utf-8")
                    raw = re.sub(r'\s+w:rsid\w+="[^"]*"', "", raw)
                    raw = re.sub(r'\s+rsid\w+="[^"]*"', "", raw)
                    data = raw.encode("utf-8")
                if item.filename == "docProps/core.xml":
                    root = etree.fromstring(data)
                    for tag in [
                        "{http://purl.org/dc/elements/1.1/}creator",
                        "{http://purl.org/dc/elements/1.1/}description",
                        "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy",
                        "{http://purl.org/dc/terms/}created",
                        "{http://purl.org/dc/terms/}modified",
                    ]:
                        el = root.find(tag)
                        if el is not None:
                            el.text = ""
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
                if item.filename == "docProps/app.xml":
                    root = etree.fromstring(data)
                    app = root.find("{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Application")
                    if app is not None:
                        app.text = ""
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def generate_docx(analysis_json, output_docx, name, project, date_str=None):
    # 读取 JSON
    with open(analysis_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    expenses = data.get("expenses", [])
    validations = data.get("validations", [])

    # 日期处理
    if date_str:
        fill_date = datetime.strptime(date_str, "%Y-%m-%d")
    else:
        fill_date = datetime.now()
    year, month, day = fill_date.year, fill_date.month, fill_date.day
    weekday = WEEKDAY_CN[fill_date.weekday()]

    # 计算合计金额
    total_amount = sum(e.get("amount", 0) for e in expenses)

    # 创建文档
    doc = Document()
    # 锁定整篇默认字体，避免 WPS / Office 渲染差异
    lock_document_font(doc)

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    # ================================================================
    # 第1部分：明细表格
    # ================================================================

    # 2.1 标题行
    make_paragraph(
        doc,
        "出差报销单明细表",
        font_size=20, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=12,
    )

    # 2.2 信息行（用 3 个段落避免溢出）
    p1 = doc.add_paragraph()
    r1 = p1.add_run(f"姓名：{name.strip()}")
    set_run_font(r1, size=12, bold=True)
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(0)
    # 不允许换行
    p1.paragraph_format.keep_together = True
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p1._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)

    p2 = doc.add_paragraph()
    project_display = project if project and project != "无" else "无"
    r2 = p2.add_run(f"项目工作号：{project_display}                    日期：{year} 年 {month} 月 {day} 日 星期 {weekday}")
    set_run_font(r2, size=12, bold=True)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)

    # 2.3 明细表格 —— 按日期排序 / 动态列宽 / 字体锁定 / 完整展示不溢出
    col_headers = ["日期", "项目", "使用地", "使用内容", "金额(RMB)"]
    num_cols = len(col_headers)

    # 拆分并整理数据行：网约车按行程单拆分为单项；误餐补贴置于末行
    main_exps = [e for e in expenses if e.get("category") != "meal_allowance"]
    meal_exps = [e for e in expenses if e.get("category") == "meal_allowance"]
    main_exps.sort(key=_sort_key)
    ordered_exps = main_exps + meal_exps

    # 构造二维矩阵（含表头），用于按内容多寡计算列宽
    matrix = [col_headers]
    for exp in ordered_exps:
        matrix.append([
            exp.get("date_display", ""),
            exp.get("item") or project,
            exp.get("city", ""),
            exp.get("description", ""),
            format_amount(exp.get("amount", 0)),
        ])
    col_widths_cm = compute_col_widths(
        matrix, total_cm=17.5, min_w=[1.8, 1.3, 1.5, 6.5, 2.2]
    )

    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    # 固定列宽布局 + 总宽，确保 WPS / Office 下列宽一致
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(17.5/2.54*1440)}" w:type="dxa"/>'))

    # 更新 tblGrid 列宽（固定布局下列宽由 gridCol 决定）
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
        table._tbl.append(grid)
    grid_cols = grid.findall(qn('w:gridCol'))
    for i, w in enumerate(col_widths_cm):
        w_twips = str(int(w / 2.54 * 1440))
        if i < len(grid_cols):
            grid_cols[i].set(qn('w:w'), w_twips)
        else:
            grid.append(parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{w_twips}"/>'))

    FONT_BODY = 12      # 正文 小四
    FONT_MIN = 10.5     # 5号，最低值

    # 表头
    hdr_row = table.rows[0]
    for i, header in enumerate(col_headers):
        cell = hdr_row.cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        set_run_font(run, size=14, bold=True)
        set_cell_width(cell, col_widths_cm[i])
        enable_word_wrap(cell)
        set_cell_vertical_alignment(cell)

    # 数据行（含误餐补贴）
    for exp in ordered_exps:
        row = table.add_row()
        values = [
            exp.get("date_display", ""),
            exp.get("item") or project,
            exp.get("city", ""),
            exp.get("description", ""),
            format_amount(exp.get("amount", 0)),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = str(val)
            p = cell.paragraphs[0]
            is_amount = (i == num_cols - 1)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_amount else WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            # 内容较多时缩小字体到 5号（最低值），保证完整展示
            uw = display_width(val)
            units_per_line = 4.72 * col_widths_cm[i]
            lines_est = math.ceil(uw / units_per_line) if units_per_line > 0 else 1
            size = FONT_BODY if lines_est <= 2 else FONT_MIN
            set_run_font(run, size=size, bold=is_amount)
            set_cell_width(cell, col_widths_cm[i])
            enable_word_wrap(cell)
            set_cell_vertical_alignment(cell)

    # 合计行
    total_row = table.add_row()
    total_labels = ["合 计", "", "", ""]
    for i, label in enumerate(total_labels):
        cell = total_row.cells[i]
        cell.text = label
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        set_run_font(run, size=14, bold=True)
        set_cell_width(cell, col_widths_cm[i])
        enable_word_wrap(cell)
        set_cell_vertical_alignment(cell)
    total_cell = total_row.cells[num_cols - 1]
    total_cell.text = format_amount(total_amount)
    p = total_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    set_run_font(run, size=14, bold=True)
    set_cell_width(total_cell, col_widths_cm[num_cols - 1])
    enable_word_wrap(total_cell)
    set_cell_vertical_alignment(total_cell)

    # 2.4 底部信息行
    make_paragraph(
        doc,
        "按照   月   日   汇率: (            )",
        font_size=14, bold=True, space_before=6,
    )
    make_paragraph(
        doc,
        "预支暂支金: (RMB)                    冲暂支金额： (RMB)",
        font_size=14, bold=True, space_before=6,
    )

    # ================================================================
    # 第2部分：附件凭证页（连续排版，跨类别不强制分页，消除空白页）
    # ================================================================

    # 按费用类型分组
    groups = {}
    for exp in expenses:
        category = exp.get("category", "其它")
        if category not in groups:
            groups[category] = []
        groups[category].append(exp)

    # 按预定义顺序排序
    sorted_categories = [c for c in GROUP_ORDER if c in groups]
    for c in sorted(groups):
        if c not in sorted_categories:
            sorted_categories.append(c)

    if sorted_categories:
        # 与明细表分页（仅一次）
        doc.add_page_break()
        attach_state = {"page_used_h": 0.0}

        from docx.oxml import OxmlElement
        LABEL_H = 0.9   # 标签行预估高度(cm)

        for category in sorted_categories:
            # 误餐补贴不放入附件页（含标签行）
            if category == "meal_allowance":
                continue
            items = groups[category]
            # 日期范围
            dates = [exp.get("date_display", "") for exp in items if exp.get("date_display")]
            date_range = ""
            if len(dates) == 1:
                date_range = dates[0]
            elif len(dates) >= 2:
                date_range = f"{dates[0]} ~ {dates[-1]}"
            # 合计金额
            cat_total = sum(exp.get("amount", 0) for exp in items)

            # 收集本类别图片（按费用项逐项收集，每项内非发票在前、发票在后）
            # 顺序：(a) 同一网约车行程的行程单+发票自然相邻，易同页；(b) 全类发票在最末。
            ordered_cat_images = []
            seen_imgs = set()
            for exp in items:
                exp_non_inv = []
                exp_inv = []
                for img_path in exp.get("page_images", []):
                    if os.path.exists(img_path) and img_path not in seen_imgs:
                        if is_invoice_image(img_path):
                            exp_inv.append(img_path)
                        else:
                            exp_non_inv.append(img_path)
                        seen_imgs.add(img_path)
                ordered_cat_images.extend(exp_non_inv)   # 行程单/截图 先
                ordered_cat_images.extend(exp_inv)        # 发票 后

            # 标注行：先预估本类首行图片高度，若当前页放不下"标签+首图"则提前换页
            if ordered_cat_images:
                try:
                    _w, _h, _ratio = image_natural_cm(ordered_cat_images[0])
                    _full = is_doc_voucher(ordered_cat_images[0]) or _ratio >= 1.1
                    if _full:
                        first_row_h = min(MAX_IMG_WIDTH_CM / _ratio, 26.7) if _ratio > 0 else 26.7
                    else:
                        # 窄图统一半行宽，按自然比例计算高度
                        first_row_h = min(_h, MAX_IMG_WIDTH_CM / 2) / _ratio if _ratio > 0 else 26.7
                except Exception:
                    first_row_h = 15.0

                if attach_state["page_used_h"] + LABEL_H + first_row_h > 26.7:
                    doc.add_page_break()
                    attach_state["page_used_h"] = 0.0
            else:
                # 无图片：标签本身放到页底则换页
                if attach_state["page_used_h"] + LABEL_H > 26.7:
                    doc.add_page_break()
                    attach_state["page_used_h"] = 0.0
            cat_cn = GROUP_CN.get(category, category)
            label = f"【{cat_cn}】{date_range} ¥{format_amount(cat_total)}"
            para = make_paragraph(doc, label, font_size=12, bold=True, space_before=2, space_after=2)
            add_shading_to_paragraph(para, "E8F0FE")
            # 标签与后续首图保持同页
            pPr = para._p.get_or_add_pPr()
            pPr.append(OxmlElement('w:keepNext'))
            attach_state["page_used_h"] += LABEL_H

            if ordered_cat_images:
                layout_images_continuous(doc, ordered_cat_images, attach_state)

    # ================================================================
    # 保存文档
    # ================================================================
    doc.save(output_docx)

    # 后处理：清空页眉页脚 + 清除元数据（尽量减少 python-docx 痕迹）
    _post_process_docx(output_docx)

    # ================================================================
    # 校验摘要输出
    filename = os.path.basename(output_docx)
    print(f"\n=== 报销单生成完成 ===")
    print(f"文件：{filename}")
    print(f"总条目：{len(expenses)}")
    print(f"合计金额：¥{format_amount(total_amount)}")

    if validations:
        print(f"\n=== 三单校验结果 ===")
        has_warning = False
        for v in validations:
            item_name = v.get("name", v.get("item", "未知"))
            status = v.get("status", v.get("result", ""))
            message = v.get("message", "")
            if status == "pass" or "齐全" in message:
                print(f"✓ {item_name}：{message or '材料齐全'}")
            else:
                print(f"⚠ {item_name}：{message or '缺少材料'}")
                has_warning = True
        if has_warning:
            print("\n⚠ 缺少材料的费用，请补充后再提交。")


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="根据分析结果 JSON 生成报销单 DOCX")
    parser.add_argument("analysis_json", help="analyze_and_validate.py 输出的 JSON 路径")
    parser.add_argument("output_docx", help="输出 DOCX 文件路径")
    parser.add_argument("--name", required=True, help="报销人姓名")
    parser.add_argument("--project", required=True, help="项目编号")
    parser.add_argument("--date", default=None, help="填表日期（yyyy-mm-dd，默认今天）")

    args = parser.parse_args()
    generate_docx(args.analysis_json, args.output_docx, args.name, args.project, args.date)
